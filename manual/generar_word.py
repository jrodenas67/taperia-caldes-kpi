from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Márgenes ───────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

NAVY   = RGBColor(0x1A, 0x2E, 0x40)
BLUE   = RGBColor(0x2E, 0x6D, 0xA4)
GREEN  = RGBColor(0x2E, 0x8A, 0x4A)
YELLOW = RGBColor(0xD4, 0xA0, 0x17)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex)
    tcPr.append(shd)

def add_section_header(doc, num, title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '1A2E40')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{num}   {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = WHITE
    doc.add_paragraph()

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)

def add_checklist(doc, items, color=None):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)

def add_info_box(doc, title, body, color_hex, title_color):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    r1 = p.add_run(title + "  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(
        int(title_color[0:2],16),
        int(title_color[2:4],16),
        int(title_color[4:6],16)
    )
    r2 = p.add_run(body)
    r2.font.size = Pt(11)
    doc.add_paragraph()

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1A2E40')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows):
        trow = tbl.rows[ri + 1]
        bg = 'D6E8F7' if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(11)
            if ci == 0:
                run.bold = True
    doc.add_paragraph()

def add_steps(doc, steps):
    for num, title, desc in steps:
        p = doc.add_paragraph()
        r1 = p.add_run(f"  {num}.  ")
        r1.bold = True
        r1.font.color.rgb = WHITE
        # workaround: just use bold text
        p = doc.add_paragraph()
        r1 = p.add_run(f"{num}. {title}")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = BLUE
        r2 = p.add_run(f"\n    {desc}")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_phrase(doc, situation, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"[{situation}]  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BLUE
    r2 = p.add_run(text)
    r2.font.size = Pt(11)

def add_ol(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(0.5)

# ══════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.rows[0].cells[0]
set_cell_bg(cell, '1A2E40')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LA TAPERIA DE CALDES")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = WHITE
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("MANUAL DE PROCEDIMIENTOS · CAMARERO/A")
r2.bold = True; r2.font.size = Pt(18); r2.font.color.rgb = RGBColor(0x7E, 0xC8, 0xF0)
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Gastrobar · v2.1 · Junio 2026")
r3.font.size = Pt(13); r3.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

doc.add_paragraph()

# pills
pills = ["⭐ Calidad de producto", "⚡ Agilidad", "🫒 Tapas & Raciones", "🤝 Trabajo en equipo", "🍺 Ambiente Gastrobar"]
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for pill in pills:
    run = p.add_run(f"  {pill}  ")
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "1", "Introducción y Objetivos")
add_body(doc, "Este manual tiene como propósito orientar al camarero/a en sus funciones diarias en La Taperia de Caldes, un gastrobar donde conviven la barra animada, las tapas de calidad y las raciones para compartir. El ambiente es informal y cercano, pero el servicio es siempre profesional y cuidado.")
add_body(doc, "En un gastrobar el ritmo es diferente al de un restaurante tradicional: el cliente puede venir solo a tomar algo, a compartir unas raciones o a cenar completo. Adaptarse a cada visita, leer al cliente y dar respuesta ágil es la clave del éxito.")
add_info_box(doc, "🍺 ¿Qué es un gastrobar?",
    "Un espacio que combina la informalidad y la energía de un bar con la calidad gastronómica de un restaurante. Tapas, raciones, bocadillos gourmet, vinos, cervezas artesanas y cócteles comparten carta. El cliente come y bebe en barra o en mesa, sin rigidez protocolaria pero con excelente producto y trato.",
    "D6E8F7", "2E6DA4")
add_info_box(doc, "📌 Cómo usar este manual",
    "Consulta este documento antes de tu primer turno y en caso de dudas operativas. Es una herramienta viva: si detectas algo mejorable, comunícalo a tu responsable.",
    "D6E8F7", "2E6DA4")

# ══════════════════════════════════════════════════════════════════════
# 2. RESPONSABILIDADES
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "2", "Identificación del Puesto y Responsabilidades")
add_body(doc, "El puesto de camarero/a ocupa un lugar estratégico dentro de la estructura organizativa del establecimiento. Su función va más allá del simple servicio de alimentos y bebidas: representa la imagen del establecimiento.")
add_table(doc,
    ["Responsabilidad", "Descripción detallada"],
    [
        ["Atención al cliente",  "Recepción cordial, toma de pedidos, atención personalizada y despedida."],
        ["Servicio de alimentos","Servir platos con la temperatura, presentación y protocolo adecuados."],
        ["Servicio de bebidas",  "Preparar y servir bebidas; conocer la carta de vinos, cócteles y cervezas."],
        ["Mise en place",        "Preparar y reponer el área de trabajo antes y durante el servicio."],
        ["Cobro",                "Gestionar la cuenta y el pago de forma precisa y agradable."],
        ["Limpieza y orden",     "Mantener la sala, barra y estaciones impecables en todo momento."],
        ["Comunicación interna", "Coordinar con cocina, barra y otros departamentos eficientemente."],
        ["Reservas",             "Gestionar solicitudes y registrar en Tableo según protocolo."],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 3. PRESENTACIÓN E HIGIENE
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "3", "Presentación e Higiene Personal")
add_body(doc, "La imagen del personal es parte inseparable de la imagen del establecimiento. Presentarse correctamente es una muestra de respeto hacia el cliente y hacia el equipo.")
add_h3(doc, "Uniforme y Apariencia")
add_checklist(doc, [
    "Uniforme limpio, planchado y en perfecto estado al inicio de cada turno.",
    "Calzado cómodo, cerrado, antideslizante y en buen estado.",
    "Cabello recogido o corto; barba aseada.",
    "Uñas cortas, limpias y sin esmalte de colores llamativos.",
    "Sin joyas (anillos, pulseras, relojes) que puedan contaminar alimentos o suponer riesgo.",
    "Perfume discreto o ninguno.",
])
add_h3(doc, "Higiene de Manos — ¿Cuándo lavarse?")
add_checklist(doc, [
    "Al inicio del turno.",
    "Después de tocar dinero, basura, cara o cabello.",
    "Antes y después de manipular alimentos.",
    "Después de usar el baño.",
])
add_info_box(doc, "⚠️ Recuerda",
    "Lava las manos con jabón antibacterial durante al menos 20 segundos. El gel hidroalcohólico complementa, pero no sustituye el lavado con agua y jabón.",
    "FFF3CD", "A07000")

# ══════════════════════════════════════════════════════════════════════
# 4. MISE EN PLACE
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "4", "Preparación y Mise en Place")
add_body(doc, "El mise en place es la base de un servicio exitoso. Una preparación meticulosa antes de que lleguen los clientes permite responder con rapidez y profesionalismo.")
add_h3(doc, "✅ Checklist de Apertura (Inicio de Turno)")
add_checklist(doc, [
    "Revisar la reserva del día en Tableo: número de cubiertos, horarios y notas especiales.",
    "Montar las mesas: mantel o salvamanteles limpio, cubiertos en posición correcta, cristalería impecable.",
    "Colocar servilletas dobladas según el estándar del establecimiento.",
    "Revisar que las sillas y taburetes de barra estén limpios y bien colocados.",
    "Preparar la estación de servicio: soporte de comandas, bolis, abrebotellas, encendedor.",
    "Asegurarse de tener suficientes cartas (menú, carta de tapas y raciones, carta de vinos y cócteles).",
    "Revisar el stock de la barra: cervezas frías, vermut, refrescos, vinos abiertos, agua.",
    "Repasar con un paño seco la cristalería de barra (cañas, copas de vino, vasos de gin).",
    "Preparar los vasos con cubiertos de postre (cucharitas, cucharas café con leche, cuchillos y tenedores).",
    "Verificar limpieza y funcionamiento de cafetera, grifo de cerveza y demás máquinas.",
    "Comprobar que la cubeta de hielo y las pinzas estén listas en la barra.",
    "Leer el parte de cocina: tapas del día, platos no disponibles, cambios de carta y alérgenos.",
    "Informarse de promociones, eventos o grupos especiales del día.",
    "Asegurarse de llevar cambio suficiente en la caja.",
])
add_h3(doc, "Montaje de Mesa — Posiciones")
add_table(doc,
    ["Elemento", "Posición / Estándar"],
    [
        ["Plato base",       "Centrado, a 2 cm del borde de la mesa"],
        ["Tenedor",          "A la izquierda del plato"],
        ["Cuchillo",         "A la derecha del plato, filo hacia adentro"],
        ["Cuchara de sopa",  "A la derecha del cuchillo"],
        ["Copa de agua",     "Arriba del cuchillo"],
        ["Copa de vino",     "Arriba y a la derecha de la copa de agua"],
        ["Servilleta",       "Sobre el plato o a la izquierda, doblada"],
        ["Pan y mantequilla","A la izquierda, con cuchillo de mantequilla"],
    ]
)
add_info_box(doc, "💡 Punto clave",
    "El orden y la limpieza no son opcionales: son requisitos esenciales. Un espacio organizado permite respuestas rápidas y reduce errores significativamente.",
    "D6E8F7", "2E6DA4")

# ══════════════════════════════════════════════════════════════════════
# 5. PROTOCOLO DE SERVICIO
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "5", "Protocolo de Servicio al Cliente")
add_body(doc, "El servicio al cliente es un arte que requiere cortesía, profesionalismo y atención al detalle. Cada interacción debe reflejar los más altos estándares de hospitalidad.")
add_h3(doc, "Flujo del Servicio — Paso a Paso")
add_steps(doc, [
    ("1",  "Recepción cálida (máx. 30 seg.)",  "Saludar con cordialidad, confirmar reserva o verificar disponibilidad. Acompañar a la mesa con elegancia. Retirar silla si procede."),
    ("2",  "Bebida de bienvenida (1-2 min)",    "Preguntar si desean algo para beber mientras ven la carta. Tomar nota y servir con prontitud."),
    ("3",  "Presentación de la carta",           "Entregar menú abierto. Presentar tapas del día, especialidades y raciones destacadas. Mencionar alérgenos si se pregunta."),
    ("4",  "Toma de pedido con precisión",       "Anotar en la comanda con letra clara: mesa, número de cubiertos, platos por posición. Confirmar alérgenos y peticiones especiales."),
    ("5",  "Servicio de tapas y raciones",       "Colocar en el centro de la mesa. Anunciar brevemente cada plato. Comprobar satisfacción 2 minutos después."),
    ("6",  "Reponer pan, bebidas y cubiertos",   "Sin que el cliente lo solicite. Vigilar vasos y cañas vacías constantemente."),
    ("7",  "Segunda oleada o platos de carta",   "Coordinar con cocina para salida simultánea. Informar si hay retraso antes de que el cliente pregunte."),
    ("8",  "Oferta de postre y café",             "Presentar carta de postres o describir opciones. Ofrecer digestivos o licores."),
    ("9",  "Gestión de la cuenta",               "Presentar cuando se solicite o sea evidente que el cliente quiere marcharse."),
    ("10", "Despedida y preparación de mesa",    "Despedir con calidez, agradecer la visita e invitar a volver. Limpiar y montar la mesa para el siguiente servicio."),
])
add_h3(doc, "Tipos de Servicio en el Gastrobar")
add_table(doc,
    ["Tipo", "Descripción y uso"],
    [
        ["Servicio de barra",          "Cliente en taburete o de pie. Trato cercano y natural. El más frecuente en La Taperia."],
        ["Servicio en mesa (tapas)",   "Platos al centro para compartir. El camarero/a indica brevemente qué es cada plato."],
        ["Servicio americano",         "Platos emplatados individualmente. Para clientes que piden carta completa."],
        ["Servicio de aperitivo",      "Vermut, caña o copa con tapa incluida. Ritmo relajado, trato personalizado con habituales."],
    ]
)
add_h3(doc, "Protocolo de Tapas y Raciones para Compartir")
add_steps(doc, [
    ("1", "Orientar la selección",    "Preguntar cuántas personas son y si quieren picar o hacer comida completa. Recomendar 3-4 tapas o 2 raciones por persona."),
    ("2", "Gestionar el ritmo",       "Acordar con cocina si quieren las tapas en oleadas o todas a la vez. Preguntar al cliente su preferencia."),
    ("3", "Platos al centro",         "Las tapas y raciones van al centro de la mesa. Aportar platos pequeños si la ración es grande."),
    ("4", "Presentar el plato",       "Al dejar cada tapa, nombrarla: \"Estas son las croquetas de jamón ibérico\" o \"La burrata con tomate cherry.\""),
    ("5", "Reponer pan y salsas",     "Con cada oleada comprobar si queda pan con tomate. Reponerlo sin que lo pidan."),
])
add_h3(doc, "Frases de Servicio Recomendadas")
phrases = [
    ("Al recibir",        "¡Buenos días/tardes! Bienvenidos. ¿Tienen reserva o prefieren que les busque mesa?"),
    ("Tomar el pedido",   "¿Se les ha ocurrido ya algo o necesitan un par de minutos más?"),
    ("Recomendar tapa",   "La especialidad de hoy es… y está siendo muy demandada."),
    ("Retirar platos",    "¿Puedo retirar esto? ¿Todo ha sido de su agrado?"),
    ("Ofrecer postre",    "¿Deseaban ver la carta de postres o tomar un café?"),
    ("Presentar cuenta",  "Aquí tiene la cuenta. Cuando quieran, sin prisa."),
    ("Despedida",         "Muchas gracias, ha sido un placer. ¡Hasta pronto!"),
]
for sit, txt in phrases:
    add_phrase(doc, sit, txt)
doc.add_paragraph()

add_h3(doc, "Gestión de Alérgenos")
add_info_box(doc, "❗ Atención",
    "Los alérgenos son una cuestión de salud y responsabilidad legal. Los 14 alérgenos principales deben estar identificados en la carta. Si el cliente pregunta por uno no identificado, consulta siempre con cocina antes de confirmar.",
    "FFF3CD", "A07000")
add_ol(doc, [
    "Preguntar siempre al inicio si algún comensal tiene alergia o intolerancia.",
    "Comunicar inmediatamente a cocina, incluso si el plato normalmente no lleva ese ingrediente.",
    "No hacer suposiciones. Confirmar con cocina antes de confirmar al cliente.",
    "Anotar la alergia en la comanda de forma visible (subrayada o en mayúsculas).",
])
add_h3(doc, "Gestión de Quejas y Reclamaciones")
add_ol(doc, [
    "Escuchar al cliente sin interrumpir. Mostrar empatía genuina.",
    "No adoptar una actitud defensiva ni discutir.",
    "Pedir disculpas con sinceridad: \"Tiene razón, lo sentimos mucho.\"",
    "Buscar una solución inmediata: cambiar el plato, ofrecer una alternativa.",
    "Comunicar al encargado si supera tu nivel de decisión.",
    "Hacer seguimiento: preguntar si la solución fue satisfactoria.",
    "Registrar la incidencia en el parte si es relevante.",
])

# ══════════════════════════════════════════════════════════════════════
# 6. SERVICIO DE BEBIDAS
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "6", "Servicio de Bebidas")
add_body(doc, "El servicio de bebidas es parte fundamental de la experiencia en el gastrobar. Un buen camarero/a conoce la carta y sabe hacer recomendaciones acertadas según el momento y el perfil del cliente.")
add_info_box(doc, "Protocolo general",
    "Servir siempre por la derecha. Bebidas frías en copa/vaso fríos; calientes en taza calentada. El agua se repone automáticamente sin que el cliente lo solicite.",
    "D6E8F7", "2E6DA4")

add_h3(doc, "Servicio de Vinos")
add_steps(doc, [
    ("1", "Mostrar la botella",  "Presentar al anfitrión con la etiqueta visible para su aprobación."),
    ("2", "Abrir la botella",    "Cortar la cápsula por debajo del cuello. Insertar el sacacorchos y extraer el corcho sin ruido."),
    ("3", "Cata previa",         "Verter una pequeña cantidad al anfitrión. Esperar su aprobación antes de servir."),
    ("4", "Servir",              "Llenar la copa hasta 2/3 de su capacidad."),
    ("5", "Dejar la botella",    "En cubo de hielo si es blanco/rosado; en la mesa con paño si es tinto."),
    ("6", "Reponer",             "Vigilar los niveles y ofrecer más cuando las copas estén casi vacías."),
])

add_h3(doc, "Tipos de Café")
add_table(doc,
    ["Tipo", "Descripción / Proporción"],
    [
        ["Café solo",       "Espresso 25 ml. Taza pequeña con cucharita y azúcar."],
        ["Café cortado",    "Espresso + pequeña cantidad de leche caliente. Taza pequeña."],
        ["Café con leche",  "Mitad café, mitad leche vaporizada. Taza grande con cuchara."],
        ["Capuchino",       "Espresso + mayor proporción de espuma. Cacao espolvoreado opcional."],
        ["Café americano",  "Espresso + agua caliente. Taza grande."],
    ]
)

add_h3(doc, "Servicio de Cerveza de Grifo")
add_ol(doc, [
    "Enjuagar el vaso con agua fría antes de servir.",
    "Inclinar el vaso 45° e iniciar el tiro desde el grifo bien abierto.",
    "Ir incorporando el vaso a la verticalidad para controlar la espuma.",
    "Dejar 1-2 cm de espuma compacta. Limpiar el exterior antes de servir.",
    "Cerrar el grifo completamente al terminar.",
])

add_h3(doc, "🍊 Servicio de Vermut")
add_info_box(doc, "El vermut: seña de identidad del gastrobar",
    "El vermut es uno de los grandes protagonistas del aperitivo. Servirlo bien — frío, con hielo, naranja o aceituna — marca la diferencia y fideliza al cliente habitual.",
    "D6E8F7", "2E6DA4")
add_table(doc,
    ["Estilo", "Copa / Vaso", "Guarnición estándar"],
    [
        ["Vermut rojo (dulce)",   "Copa balón o vaso de vermut con mucho hielo", "Rodaja de naranja + aceituna manzanilla"],
        ["Vermut blanco (dry)",   "Copa balón con hielo",                         "Piel de limón + aceituna"],
        ["Vermut de grifo",       "Caña de vermut o chato",                        "Preguntar siempre: \"¿Con naranja o solo?\""],
    ]
)
add_checklist(doc, [
    "Servir siempre muy frío, con abundante hielo y sin ahorrarlo.",
    "Ofrecer tapa al vermut si el establecimiento lo incluye (aceitunas, boquerones, patatas).",
    "Al vermut de grifo: limpiar el grifo antes del primer servicio del día.",
])

add_h3(doc, "Servicio de Gin-Tonic y Cócteles")
add_table(doc,
    ["Elemento", "Estándar de servicio"],
    [
        ["Copa",     "Copa balón grande (globo). Nunca un vaso recto."],
        ["Hielo",    "Abundante. Llenar la copa hasta arriba antes de añadir la ginebra."],
        ["Ginebra",  "Verter despacio sobre el hielo. Nunca agitar."],
        ["Tónica",   "Verter en vertical desde arriba, lentamente, para conservar el gas."],
        ["Guarnición","Según la ginebra. Mínimo: ralladura de limón o pepino. Presentar con pajita y posavasos."],
    ]
)
add_info_box(doc, "⚠️ Gin-tonic de autor",
    "Si en carta hay gin-tonics con guarnición específica (bayas de enebro, cardamomo, romero…), respetar la receta exacta. La presentación es parte del producto.",
    "FFF3CD", "A07000")

add_h3(doc, "Cervezas Artesanas y Especiales")
add_table(doc,
    ["Estilo", "Descripción breve", "Maridaje con tapas"],
    [
        ["Lager / Pilsner", "Suave, refrescante, poco amarga. La más pedida.",             "Croquetas, frituras, bocadillos"],
        ["Weizen (de trigo)","Afrutada, cremosa, con notas de plátano.",                   "Embutidos, quesos suaves"],
        ["IPA / Pale Ale",   "Aromática, con más amargor de lúpulo.",                      "Carnes especiadas, quesos curados"],
        ["Stout / Porter",   "Oscura, tostada, sabor a café o chocolate.",                 "Morcilla, patés, postres de chocolate"],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 7. UPSELLING Y CROSS-SELLING
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "7", "Técnicas de Venta: Upselling y Cross-selling")
add_body(doc, "El upselling bien aplicado mejora la experiencia del cliente y aumenta la rentabilidad. No se trata de presionar, sino de sugerir con conocimiento y autenticidad.")
add_info_box(doc, "💡 Regla de oro",
    "Recomienda lo que tú probarías. La autenticidad convence más que cualquier técnica de venta.",
    "E8F4E8", "2E8A4A")
add_h3(doc, "Principios del Upselling Efectivo")
add_checklist(doc, [
    "Conoce la carta a fondo: ingredientes, técnicas y pairings recomendados.",
    "Haz sugerencias en el momento adecuado, no de forma repetitiva.",
    "Usa lenguaje sensorial: \"es muy sabroso\", \"tiene mucho éxito\", \"va perfecto con…\"",
    "Adapta la recomendación al perfil del cliente.",
])
add_h3(doc, "Oportunidades por Fase del Servicio")
add_table(doc,
    ["Momento", "Sugerencia de upselling"],
    [
        ["Al llegar",       "¿Desearían una bebida de bienvenida? (vermut, cava, agua con gas)"],
        ["Con la carta",    "Mencionar las especialidades de precio superior o la tabla de quesos."],
        ["Al pedir bebida", "\"Este plato va muy bien con nuestro blanco de la casa.\""],
        ["Entre tapas",     "Ofrecer pan artesano con tomate, extras o guarnición adicional."],
        ["Tras las raciones","\"¿Nos dejan tentarles con un postre? El fondant de chocolate es impresionante.\""],
        ["Con el café",     "Sugerir un digestivo, licor o infusión especial."],
    ]
)
add_h3(doc, "Cross-selling: Vender en Paralelo")
add_info_box(doc, "🔄 Diferencia clave",
    "Upselling = versión mejor del mismo producto. Cross-selling = producto diferente que complementa al anterior.",
    "D6E8F7", "2E6DA4")
add_table(doc,
    ["El cliente pide…", "Sugerencia de cross-selling"],
    [
        ["Cualquier tapa o plato", "\"¿Le ponemos una cestita de pan con tomate?\""],
        ["Cerveza de caña",        "\"¿Lo acompaña con unas aceitunas o patatas bravas?\""],
        ["Vermut",                 "\"¿Le pongo unas boquerones o aceitunas para acompañar?\""],
        ["Plato principal",        "\"¿Quiere añadir una guarnición de patatas o ensalada?\""],
        ["Postre",                 "\"¿Un café o una copita de cava para terminar?\""],
        ["Café solo",              "\"¿Con una galletita o un chupito de licor?\""],
    ]
)
add_h3(doc, "Frases que Funcionan")
add_checklist(doc, [
    "\"Muchos clientes lo piden con… ¿se lo añado?\"",
    "\"Va muy bien acompañado de… ¿lo probamos?\"",
    "\"Para compartir, también tenemos… está muy bueno.\"",
    "\"¿Pongo algo para picar mientras esperan?\"",
])
add_info_box(doc, "⚠️ Importante",
    "Una sola sugerencia por momento. Si el cliente dice que no, no insistir. El objetivo es facilitar, no presionar.",
    "FFF3CD", "A07000")

# ══════════════════════════════════════════════════════════════════════
# 8. GESTIÓN DEL COBRO
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "8", "Gestión del Cobro")
add_body(doc, "El cobro es el último punto de contacto con el cliente antes de la despedida. Una gestión ágil y sin errores cierra la experiencia de forma positiva.")
add_steps(doc, [
    ("1", "Preparar la cuenta",  "Revisar la comanda y la factura antes de presentarla. Verificar que no falte ni sobre ningún elemento."),
    ("2", "Presentar la cuenta", "Entregar en cartera o bandeja, por la derecha del anfitrión. No se entrega sin que el cliente lo pida o sea evidente que desea marcharse."),
    ("3", "Pago en efectivo",    "Contar el dinero a la vista del cliente. Indicar el cambio en voz alta antes de devolverlo."),
    ("4", "Pago con tarjeta",    "Acercar el datáfono al cliente. Esperar confirmación de la operación antes de dar el pago por completado."),
    ("5", "Factura",             "Preguntar si necesitan factura. Tomar los datos fiscales y tramitarla correctamente."),
    ("6", "Despedida",           "Agradecer, despedirse por el nombre si lo conoces e invitar a volver."),
])
add_info_box(doc, "⚠️ Nunca",
    "Nunca insinuar, reclamar ni comentar sobre la propina. Jamás hacer esperar al cliente más de 5 minutos con la cuenta presentada. Si hay discrepancia, la duda beneficia siempre al cliente.",
    "FFF3CD", "A07000")

# ══════════════════════════════════════════════════════════════════════
# 9. MAQUINARIA
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "9", "Uso y Mantenimiento de Maquinaria")
add_body(doc, "El conocimiento profundo de la maquinaria es esencial para la operación diaria. Cada equipo requiere procedimientos específicos de uso y mantenimiento.")
add_table(doc,
    ["Equipo", "Puntos clave de uso y mantenimiento"],
    [
        ["Cafetera profesional", "Purgar antes del primer café. Limpiar portafiltros y grupos al terminar el turno. Descalcificar según frecuencia indicada."],
        ["Molinillo de café",    "Ajustar la molienda según extracción. Limpiar tolva y molinillo semanalmente."],
        ["Grifo de cerveza",     "Limpiar líneas con agua al inicio. Purgado y limpieza de válvulas según protocolo."],
        ["Lavavasos industrial", "Verificar niveles de detergente y abrillantador diariamente. Limpiar filtros al final del turno."],
        ["TPV / Datáfono",       "Verificar conexión y papel al inicio. Reportar errores de inmediato."],
        ["Frigorífico y cámaras","Controlar temperatura (2-8°C). Registrar incidencias. No sobrecargar."],
    ]
)

# ══════════════════════════════════════════════════════════════════════
# 10. DURANTE EL SERVICIO
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "10", "Tareas Durante el Servicio")
add_body(doc, "La atención a las mesas debe ser constante y proactiva. El cliente nunca debe tener que buscar al camarero/a con la mirada para necesidades básicas.")
add_h3(doc, "Tiempos Orientativos — Gastrobar")
add_table(doc,
    ["Momento", "Tiempo objetivo"],
    [
        ["Bienvenida / barra",          "30 segundos"],
        ["Primera bebida",               "2 minutos"],
        ["Primera tapa / ración",        "8-12 minutos"],
        ["Tapas siguientes (oleadas)",   "5-8 minutos"],
        ["Cuenta solicitada",            "2 minutos"],
    ]
)
add_info_box(doc, "⏱️ Ritmo de gastrobar",
    "El cliente de gastrobar tolera menos espera que en un restaurante formal. Si una tapa tarda más de 15 minutos, hay que informar proactivamente y ofrecer algo mientras espera (pan con tomate, aceitunas).",
    "FFF3CD", "A07000")
add_h3(doc, "Monitoreo Continuo")
add_checklist(doc, [
    "Realizar una pasada visual de todas las mesas cada 5-8 minutos.",
    "Reponer agua, pan y cubiertos sin que el cliente lo solicite.",
    "Retirar vajilla sucia cuando el cliente ha terminado (nunca antes de que todos hayan acabado).",
    "Atender señales no verbales: cliente mirando alrededor, vaso vacío, platos apilados.",
])
add_h3(doc, "Coordinación con Cocina y Barra")
add_checklist(doc, [
    "Pasar los pedidos con claridad: mesa, posición, alergias y observaciones.",
    "Confirmar tiempos de tapas al pasar la comanda, especialmente en mesas grandes.",
    "Si hay retraso, informar al cliente proactivamente antes de que pregunte.",
    "No llevar tapas a la mesa si no están todos los platos de la misma oleada listos.",
])

# ══════════════════════════════════════════════════════════════════════
# 11. RESERVAS
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "11", "Protocolo de Reservas")
add_body(doc, "Las reservas son la primera interacción del cliente con el establecimiento. Gestionarlas correctamente sienta las bases de una buena experiencia.")
add_h3(doc, "Canales de Reserva")
add_table(doc,
    ["Canal", "Gestión"],
    [
        ["📸 Instagram",            "El cliente solicita reserva por mensaje directo en el perfil oficial."],
        ["🎵 TikTok",               "Reservas por mensaje directo o enlace en la bio del perfil."],
        ["🔍 Google Business",      "Solicitud directamente desde la ficha de Google."],
        ["📞 Telefónica / presencial","El cliente llama o se acerca al establecimiento directamente."],
    ]
)
add_h3(doc, "Proceso de Gestión en Tableo")
add_steps(doc, [
    ("1", "Recibir la solicitud",   "El cliente indica fecha, hora, número de comensales y necesidades especiales (cumpleaños, alergia, silla para bebé…)."),
    ("2", "Registrar en Tableo",    "Introducir todos los datos: nombre, teléfono/email, fecha, hora, número de personas y observaciones."),
    ("3", "Confirmación automática","El cliente recibe confirmación con las condiciones de la reserva. Verificar que se haya enviado correctamente."),
    ("4", "Día de la reserva",      "Recibir al cliente con calidez, confirmar nombre y asignar mesa según disponibilidad en ese momento."),
])
add_info_box(doc, "📌 Normas clave de reservas",
    "Máximo 6 personas por reserva. En ningún caso se puede garantizar ni asegurar una mesa específica al cliente. La asignación de mesa se realiza el día de la reserva según disponibilidad.",
    "FFF3CD", "A07000")

# ══════════════════════════════════════════════════════════════════════
# 12. HIGIENE Y SEGURIDAD
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "12", "Normas de Higiene y Seguridad Alimentaria")
add_body(doc, "La higiene y seguridad son pilares fundamentales en la industria gastronómica. Su incumplimiento puede tener consecuencias legales y de salud pública.")
add_h3(doc, "Manejo Seguro de Alimentos")
add_table(doc,
    ["Principio", "Aplicación práctica"],
    [
        ["Temperatura de conservación","Frío: 0-5°C. Caliente: >65°C. Zona de peligro: 5-65°C (evitar siempre)."],
        ["Contaminación cruzada",       "Utensilios separados para crudos y cocinados. Nunca mezclar."],
        ["Rotación FIFO",               "Primero en entrar, primero en salir. Etiquetar con fecha de apertura."],
        ["Vida útil",                   "Respetar fechas de caducidad. Desechar productos dudosos."],
    ]
)
add_h3(doc, "Protocolos de Emergencia")
add_checklist(doc, [
    "Conocer la ubicación exacta de los extintores, salidas de emergencia y botiquín.",
    "En caso de accidente o lesión: atender al afectado, avisar al responsable, registrar el incidente.",
    "En caso de incendio: avisar, evacuar, no usar ascensores.",
    "Números de emergencia visibles y accesibles en la estación de trabajo.",
    "Conocer el protocolo de actuación ante reacción alérgica grave (anafilaxia).",
])

# ══════════════════════════════════════════════════════════════════════
# 13. CIERRE Y LIMPIEZA
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "13", "Procedimientos de Cierre y Limpieza")
add_body(doc, "Un cierre meticuloso garantiza que el siguiente turno comience en óptimas condiciones y mantiene los estándares de higiene del establecimiento.")
add_h3(doc, "Limpieza de Barra y Estaciones")
add_checklist(doc, [
    "Superficies de barra desinfectadas completamente.",
    "Cafetera y molinillo limpios (grupos, portafiltros, vaporizador).",
    "Grifo de cerveza purgado y limpiado. Grifo de vermut limpiado si corresponde.",
    "Lavavasos: filtros limpios, niveles de detergente y abrillantador verificados.",
    "Derrames y manchas eliminados de todas las superficies.",
    "Cubeta de hielo vaciada y limpiada al cierre.",
])
add_h3(doc, "Sala y Mobiliario")
add_checklist(doc, [
    "Todas las mesas limpias, sin migas ni manchas.",
    "Sillas y taburetes colocados correctamente y revisados.",
    "Suelo barrido/fregado en el área de responsabilidad.",
    "Decoración y elementos de mesa en buen estado.",
])
add_h3(doc, "Utensilios y Cristalería")
add_checklist(doc, [
    "Toda la cristalería lavada, secada y almacenada en su lugar.",
    "Copas de gin, cañas y vasos de vermut repasados con paño seco.",
    "Cubertería clasificada y guardada por tipo.",
    "Inventario visual de faltantes o roturas (registrar).",
])
add_h3(doc, "🔑 Preparación para el Siguiente Turno")
add_checklist(doc, [
    "Vasos rellenados con cucharitas de café.",
    "Vasos rellenados con cucharas de café con leche.",
    "Vasos rellenados con cuchillos de postre.",
    "Vasos rellenados con tenedores de postre.",
    "Servilletas, manteles y paños repuestos.",
    "Condimentos y complementos verificados y rellenados.",
    "Stock mínimo asegurado para el siguiente turno.",
    "Cartas de menú limpias y en su lugar.",
])
add_h3(doc, "Reporte Final")
add_checklist(doc, [
    "Incidencias documentadas por escrito en el parte.",
    "Faltantes o averías comunicados al responsable.",
    "Información relevante transmitida al turno siguiente.",
    "Checklist firmado y entregado al encargado.",
])
add_info_box(doc, "🔑 Recordatorio de cierre clave",
    "La preparación de los vasos con cubiertos de postre (cucharitas, cucharas de café con leche, cuchillos y tenedores) es un paso crítico que agiliza enormemente el inicio del siguiente turno.",
    "E8F4E8", "2E8A4A")

# ══════════════════════════════════════════════════════════════════════
# 14. CONCLUSIÓN
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "14", "Conclusión y Buenas Prácticas")
add_body(doc, "El éxito del servicio depende del compromiso individual con la excelencia. Seguir los procedimientos de este manual es una herramienta para el crecimiento profesional y la satisfacción personal en el trabajo.")
add_table(doc,
    ["Valor", "Descripción"],
    [
        ["😊 Actitud positiva",   "Una sonrisa genuina y disposición servicial transforman la experiencia del cliente."],
        ["📚 Formación continua", "El aprendizaje constante abre nuevas oportunidades y mantiene la motivación."],
        ["🔍 Atención al detalle","Los pequeños detalles marcan la diferencia entre un servicio correcto y uno excepcional."],
        ["🤝 Trabajo en equipo",  "El mejor servicio individual fracasa sin coordinación con el equipo."],
        ["⚡ Proactividad",       "Anticiparse a las necesidades del cliente antes de que las exprese."],
        ["💬 Comunicación",       "Informar siempre a compañeros y superiores de incidencias y novedades."],
    ]
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"La excelencia no es un acto, sino un hábito. Cada interacción con el cliente es una oportunidad para demostrar profesionalismo y crear experiencias memorables."')
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = NAVY

# ══════════════════════════════════════════════════════════════════════
# GLOSARIO
# ══════════════════════════════════════════════════════════════════════
add_section_header(doc, "G", "Glosario de Términos")
glosario = [
    ("Mise en place",     "Preparación y disposición de todos los elementos necesarios antes del servicio."),
    ("Comanda",           "Documento (físico o digital) donde se anotan los pedidos de los clientes."),
    ("Tapa",              "Pequeña porción de comida que acompaña a una bebida o se sirve como aperitivo."),
    ("Ración",            "Porción más grande pensada para compartir entre 2-3 personas. Formato característico del gastrobar."),
    ("Media ración",      "Mitad de una ración. Permite probar más platos sin excederse en cantidad."),
    ("Oleada",            "Salida de varias tapas o raciones a la vez para una misma mesa."),
    ("Vermut",            "Vino aromatizado con hierbas y especias. Bebida aperitivo. Se sirve frío con hielo y guarnición."),
    ("Gin-tonic",         "Combinado de ginebra y agua tónica en copa balón con abundante hielo y guarnición aromática."),
    ("Caña",              "Vaso pequeño de cerveza de grifo (aprox. 200 ml). La medida más habitual en barra."),
    ("Tubo",              "Vaso alargado de cerveza de grifo de mayor capacidad que la caña (aprox. 300-330 ml)."),
    ("Aperitivo",         "Momento previo a la comida donde se toman bebidas y pequeños snacks. Hora de máximo consumo de vermut."),
    ("Upselling",         "Técnica de venta que consiste en sugerir la versión mejorada o de mayor valor de lo que el cliente ya ha pedido."),
    ("Cross-selling",     "Técnica de venta cruzada: sugerir productos complementarios al pedido del cliente."),
    ("FIFO",              "First In, First Out: sistema de rotación de stock (primero en entrar, primero en salir)."),
    ("Alérgeno",          "Sustancia que puede causar reacción alérgica. Existen 14 alérgenos de declaración obligatoria."),
    ("Pase",              "Salida simultánea de platos del mismo tiempo para una misma mesa."),
    ("Maridaje",          "Combinación armónica de vino (u otra bebida) con un plato determinado."),
    ("Tableo",            "Plataforma digital de gestión de reservas usada en este establecimiento."),
    ("EPI",               "Equipo de Protección Individual (guantes, delantal, etc.)."),
    ("Anfitrión",         "Persona que organiza o paga la mesa; a quien se presenta la cuenta y se sirve primero en la cata."),
    ("TPV",               "Terminal Punto de Venta. Sistema de cobro electrónico / caja registradora."),
]
add_table(doc,
    ["Término", "Definición"],
    [(k, v) for k, v in glosario]
)

# ══════════════════════════════════════════════════════════════════════
# PIE
# ══════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph("Manual de Procedimientos — Camarero/a · v2.1 Gastrobar · Junio 2026 · La Taperia de Caldes")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

output = "/Users/romaglobal/taperia-caldes-kpi/manual/Manual_Camarero_LaTaperia_v2.1.docx"
doc.save(output)
print(f"✅ Guardado en: {output}")
